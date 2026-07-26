from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "Document.docx"
BACKUP = ROOT / "Document.pre-structure-backup.docx"
OUTPUT = ROOT / "Document.docx"

INDIGO = RGBColor(79, 70, 229)
NAVY = RGBColor(17, 24, 39)
SLATE = RGBColor(71, 85, 105)
LIGHT = "E2E8F0"


def source_paragraphs(path: Path) -> list[str]:
    document = Document(path)
    return [p.text.strip() for p in document.paragraphs if p.text.strip()]


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def add_label_value(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.color.rgb = NAVY
    value_run = paragraph.add_run(value)
    value_run.font.color.rgb = SLATE


def add_bullet(document: Document, text: str, level: int = 0) -> None:
    clean = re.sub(r"^[•\-]\s*", "", text).strip()
    paragraph = document.add_paragraph(clean, style="List Bullet" if level == 0 else "List Bullet 2")
    paragraph.paragraph_format.space_after = Pt(3)


def add_question(document: Document, question: str, answer: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(question)
    run.bold = True
    run.font.color.rgb = INDIGO
    answer_paragraph = document.add_paragraph(answer)
    answer_paragraph.paragraph_format.space_after = Pt(4)


def add_rule(document: Document) -> None:
    paragraph = document.add_paragraph()
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), LIGHT)
    borders.append(bottom)
    properties.append(borders)


def clean_text(text: str) -> str:
    replacements = {
        "Current Address - Hyderabad": "Current Address - Bangalore",
        "Highradius Coporation": "HighRadius Corporation",
        "Highradius Technology": "HighRadius Technology",
        "LNR Risk LexusNexus": "LNR Risk — LexisNexis",
        "Axis Maxlife": "Axis Max Life",
        "whatsapp": "WhatsApp",
        "Github": "GitHub",
        "CoPilot": "Copilot",
        "Besic": "Basic",
        "roadmaps,lead": "roadmaps, leading",
        "enterprise platform": "enterprise platforms",
        "chain-of-thought execution": "multi-step reasoning and workflow orchestration",
        "META based infrastructure": "Meta-based infrastructure",
        "META internal": "Meta internal",
        "Client - META": "Client - Meta",
    }
    cleaned = text
    for old, new in replacements.items():
        cleaned = cleaned.replace(old, new)
    return cleaned.strip()


def add_experience(document: Document, paragraphs: list[str]) -> None:
    document.add_heading("Professional Experience", level=1)
    previous_normalized = ""
    for index, raw in enumerate(paragraphs):
        text = clean_text(raw)
        if index == 0 or text in {"Professional Experience -", "Company Details :"}:
            continue
        next_text = clean_text(paragraphs[index + 1]) if index + 1 < len(paragraphs) else ""
        is_role = bool(re.match(r"^\(.+\)$", next_text)) or (
            any(title in text for title in ("Data Scientist", "Solution Architect", "Technical Assistant"))
            and not text.startswith(("Developed", "Designed", "Led ", "Applied", "Managed", "Mentored"))
        )
        if is_role:
            document.add_heading(text.replace(" --> ", " — "), level=2)
        elif re.match(r"^\(.+\)$", text):
            paragraph = document.add_paragraph(text.strip("() "))
            paragraph.runs[0].italic = True
            paragraph.runs[0].font.color.rgb = SLATE
        normalized = re.sub(r"[\s.]+$", "", text).lower()
        if not is_role and not re.match(r"^\(.+\)$", text) and normalized != previous_normalized:
            add_bullet(document, text)
        previous_normalized = normalized


def add_projects(document: Document, paragraphs: list[str]) -> None:
    document.add_heading("Selected Professional Projects", level=1)
    company = ""
    for raw in paragraphs:
        text = clean_text(raw)
        if text == "Professional Experience Project-":
            continue
        if text.endswith(("–", "-")) and len(text.split()) <= 5:
            company = text.rstrip("–- ").strip()
            document.add_heading(company, level=2)
            continue
        if " - " in text:
            title, description = text.split(" - ", 1)
            document.add_heading(title.strip(), level=3)
            document.add_paragraph(description.strip())
        else:
            add_bullet(document, text)


def add_section_from_range(
    document: Document,
    heading: str,
    paragraphs: list[str],
    skip_markers: set[str],
) -> None:
    document.add_heading(heading, level=1)
    for raw in paragraphs:
        text = clean_text(raw)
        if text in skip_markers:
            continue
        if text.endswith(":") and len(text.split()) <= 5:
            document.add_heading(text.lstrip("•- ").rstrip(":"), level=2)
        else:
            add_bullet(document, text)


def build_document(source: list[str]) -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.font.color.rgb = SLATE
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size in (("Title", 27), ("Heading 1", 17), ("Heading 2", 13), ("Heading 3", 11)):
        style = document.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = INDIGO if style_name != "Title" else NAVY

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Chandramouli Das")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Professional Knowledge Base · Recruiter & Portfolio Edition")
    subtitle_run.bold = True
    subtitle_run.font.color.rgb = INDIGO
    add_rule(document)

    document.add_heading("Professional Quick Reference", level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = "Light Shading Accent 1"
    header = table.rows[0].cells
    header[0].text = "Topic"
    header[1].text = "Verified information"
    set_repeat_table_header(table.rows[0])
    for cell in header:
        set_cell_shading(cell, "4F46E5")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    quick_reference = [
        ("Current role", "Lead Data Scientist at Wipro, working with Meta"),
        ("Current location", "Bangalore, India"),
        ("Experience", "Over 7 years across Data Science, Generative AI, AI/ML architecture, and automation"),
        ("Notice period", "Officially 60 days; negotiable to 30 days"),
        ("Core expertise", "Generative AI, Data Science, AI Leadership, Agentic AI, RAG systems, NLP, and intelligent document automation"),
        ("Leadership", "Directly led teams at Gramener/Straive and Proxima; contributes to delivery leadership and technical direction at Wipro"),
        ("Mentoring", "Mentored 1,000+ learners across multiple platforms"),
        ("Appointment", "https://topmate.io/chandramouli_das"),
    ]
    for label, value in quick_reference:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value

    document.add_heading("Contact & Professional Links", level=1)
    add_label_value(document, "Mobile", "+91 9674078742 (call or WhatsApp)")
    add_label_value(document, "Home base", "Kolkata, India")
    add_label_value(document, "Current location", "Bangalore, India")
    add_label_value(document, "LinkedIn", "https://www.linkedin.com/in/chandramouli-das-38a7921a5/")
    add_label_value(document, "GitHub", "https://github.com/Chandramouli-Das")
    add_label_value(document, "Appointment booking", "https://topmate.io/chandramouli_das")
    add_label_value(document, "Notice period", "60 days officially; negotiable to 30 days")

    document.add_heading("Professional Profile", level=1)
    for text in source[10:21]:
        clean = clean_text(text)
        if clean == "Core strengths include:":
            document.add_heading("Core Strengths", level=2)
        elif clean.startswith("•"):
            add_bullet(document, clean)
        else:
            document.add_paragraph(clean)

    document.add_heading("Recruiter FAQ", level=1)
    faq = [
        ("What is Chandramouli’s current role?", "Lead Data Scientist at Wipro, currently working on Meta-focused Generative AI infrastructure, internal tooling, and automation initiatives."),
        ("How much professional experience does he have?", "Over 7 years of full-time experience spanning Data Science, Generative AI, AI/ML solution architecture, financial technology, legal risk, HR technology, and business automation."),
        ("Where is he currently based?", "Bangalore, India. His home base is Kolkata."),
        ("What is his notice period?", "The official notice period is 60 days and is negotiable to 30 days."),
        ("Which roles best match his background?", "Lead Data Scientist, Generative AI Lead, AI/ML Solution Architect, Data Science leadership, and senior roles involving agentic AI, RAG, NLP, or intelligent automation."),
        ("What are his strongest technical areas?", "Generative AI, LLM application design, retrieval-augmented generation, LangGraph agentic workflows, text-to-SQL, OCR and document intelligence, NLP, predictive modeling, FastAPI, vector databases, and cloud-based AI delivery."),
        ("Has he led teams?", "Yes. At Gramener/Straive, he directly led a five-member Generative AI delivery team. At Proxima Systems, he led a cross-functional AI/ML team through the full delivery lifecycle. At Wipro, he contributes as part of the leadership group for Meta-focused GenAI initiatives through technical direction, architecture ownership, prioritization, and stakeholder alignment, without direct people-management responsibility. He has also led larger delivery and roadmap activities involving 10+ members earlier in his career."),
        ("Which projects best demonstrate his experience?", "Notable work includes an agentic SQL chatbot for Axis Max Life, an OCR and GenAI legal-document pipeline for LexisNexis-related risk workflows, Logitech Sherlock text-to-SQL prompt enhancement, HR document assistants, and financial-risk prediction products at HighRadius."),
        ("Which industries has he worked in?", "Financial technology, insurance, HR technology, legal and risk, enterprise automation, education, and research."),
        ("What is his education?", "M.Tech in Data Science from KIIT with a 9.01 CGPA, a postgraduate certification in Machine Learning and Deep Learning from IIIT Bangalore, and a B.Tech in Computer Science and Engineering."),
        ("Does he have publications or patents?", "Yes. His profile includes a filed CAC Optimiser patent and peer-reviewed publications covering recommender systems, blockchain-supported supply chains, medical AI, and audio visualization."),
        ("Does he have mentoring experience?", "Yes. He has taught or mentored more than 1,000 learners and supported capstone projects, interviews, résumé reviews, and career transitions."),
        ("How can someone schedule a conversation?", "Use his Topmate booking page at https://topmate.io/chandramouli_das or contact him through LinkedIn."),
        ("What are his compensation expectations?", "Not specified in this knowledge base. Discuss compensation directly with Chandramouli."),
        ("What work arrangement does he prefer?", "Remote, hybrid, or on-site preference is not specified. Confirm directly for the specific opportunity."),
        ("Is he open to relocation?", "Relocation preference is not specified. Confirm directly with Chandramouli."),
        ("Why is he considering a change?", "The reason for exploring opportunities is not specified and should not be inferred. Ask him directly."),
        ("What is his work authorization or visa status?", "Not specified in this knowledge base. Confirm directly for the relevant country and role."),
    ]
    for question, answer in faq:
        add_question(document, question, answer)

    add_experience(document, source[28:60])
    document.add_heading("Leadership Profile", level=1)
    leadership_profile = [
        "Leadership experience spans Proxima Systems, Gramener/Straive, and Wipro, combining direct team management with technical leadership, delivery ownership, stakeholder influence, and hands-on AI architecture.",
        "At Gramener/Straive, led a five-member team delivering Generative AI solutions, coordinating priorities, solution design, execution, client communication, and end-to-end delivery.",
        "At Proxima Systems, led a cross-functional AI/ML team through requirements, architecture, development, mentoring, deployment, and enterprise integration.",
        "At Wipro, contributes as part of the leadership group for Meta-focused GenAI initiatives. Although the role does not include direct people management, it involves technical direction, architecture ownership, use-case prioritization, stakeholder alignment, and influencing delivery decisions across product and engineering teams.",
        "Leadership strengths include translating business problems into AI roadmaps, creating clarity across cross-functional teams, balancing hands-on engineering with delivery governance, mentoring practitioners, and communicating effectively with clients and senior stakeholders.",
    ]
    for statement in leadership_profile:
        add_bullet(document, statement)
    add_projects(document, source[60:76])
    add_section_from_range(document, "Leadership, Volunteering & Training", source[76:85], {"Volunteer Work -"})
    add_section_from_range(document, "Education, Research & Publications", source[85:112], {"Education -"})
    add_section_from_range(
        document,
        "Professional Certifications",
        source[112:157],
        {"Certifications (Long Term) -", "Certification (Short Term) -"},
    )
    add_section_from_range(document, "Personal AI & Data Projects", source[157:167], {"Personal Project -"})

    document.add_heading("Technical Capabilities", level=1)
    capability_groups = {
        "Generative AI & LLM Systems": "Generative AI, LLM application architecture, prompt engineering, RAG, agentic AI, LangChain, LangGraph, LlamaIndex, function calling, LLM fine-tuning, LoRA, PEFT, Hugging Face Transformers, and multimodal AI.",
        "Data Science & Machine Learning": "Predictive modeling, regression, classification, clustering, ensemble learning, XGBoost, LightGBM, anomaly detection, feature engineering, optimization, model explainability, and experimentation.",
        "NLP, Documents & Vision": "Natural language processing, text classification, named entity recognition, document understanding, summarization, OCR, computer vision, image classification, object detection, and multimodal processing.",
        "AI Engineering & Delivery": "Python, FastAPI, REST APIs, vector databases including FAISS and Pinecone, Docker, Kubernetes, CI/CD, model serving, monitoring, data pipelines, and microservice architecture.",
        "Cloud & Data Platforms": "Azure, AWS, GCP, SageMaker, Azure ML, Vertex AI, Redshift, Athena, PostgreSQL, S3, distributed computing, Spark, and streaming technologies.",
        "Leadership & Product": "AI product strategy, technical roadmaps, cross-functional team leadership, Agile and Scrum, JIRA, stakeholder communication, client engagement, pre-sales consulting, mentoring, and responsible AI.",
    }
    for heading, description in capability_groups.items():
        document.add_heading(heading, level=2)
        document.add_paragraph(description)

    document.add_heading("Tools & Platforms", level=2)
    document.add_paragraph(clean_text(source[170]))
    document.add_heading("Domain Expertise", level=2)
    document.add_paragraph(clean_text(source[172]))
    document.add_heading("Languages", level=2)
    document.add_paragraph("English, Bengali, Hindi")

    add_section_from_range(
        document,
        "Teaching, Mentorship & Career Guidance",
        source[174:215],
        {"Additional -"},
    )
    add_section_from_range(
        document,
        "Personal Interests & Conversation Starters",
        source[217:],
        {"⸻", "🧍‍♂️ Personal Interests & Fun Facts"},
    )

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("Knowledge Base Usage Note", level=1)
    document.add_paragraph(
        "This document is the factual source for MoonGPT. The assistant should answer "
        "from this content, avoid inventing missing recruiter information, and direct "
        "users to Chandramouli when compensation, relocation, work arrangement, reason "
        "for change, or work-authorization details are requested."
    )
    return document


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    if not BACKUP.exists():
        shutil.copy2(SOURCE, BACKUP)
    source = source_paragraphs(BACKUP)
    document = build_document(source)
    document.core_properties.title = "Chandramouli Das — Professional Knowledge Base"
    document.core_properties.subject = "Structured source for recruiter questions and MoonGPT"
    document.core_properties.author = "Chandramouli Das"
    document.save(OUTPUT)
    print(f"Structured document written to {OUTPUT}")
    print(f"Original preserved at {BACKUP}")


if __name__ == "__main__":
    main()
