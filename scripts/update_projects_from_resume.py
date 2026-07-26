from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCUMENT_PATH = ROOT / "Document.docx"
BACKUP_PATH = ROOT / "Document.before-resume-project-update.docx"


WIPRO_PROJECTS = [
    (
        "AWB Runbook Automation — Nexus Tools",
        "Technologies: Llama 4, Nexus Runbook, Agentic Workflow Builder, "
        "Sapling, MetaClaw, MCP-compatible tools.",
        "Leading the design and delivery of an enterprise agentic automation "
        "platform for converting operational runbooks into executable AI "
        "workflows. The system interprets a task, selects the appropriate "
        "MCP-compatible tools, coordinates specialist agents, and executes "
        "multi-step actions with validation and failure-handling checkpoints. "
        "Built on Meta Nexus Runbook, Agentic Workflow Builder, Sapling, and "
        "MetaClaw, it reduces repetitive operational effort and provides a "
        "reusable foundation for scalable business-process automation.",
    ),
    (
        "People Data Review — Agentic Execution",
        "Technologies: GPT-4.1 Mini, Agentic AI, tool orchestration, parallel "
        "multi-agent workflows, human-in-the-loop validation.",
        "Architected and delivered an agentic review system for completing "
        "more than 160 compliance and people-risk assessment questions from "
        "enterprise source material. Parallel specialist agents retrieve "
        "section-specific evidence, generate grounded answers, cross-check "
        "results, and route uncertain cases for human validation. Section-aware "
        "reasoning reduced hallucinations by 70%, achieved 99% groundedness "
        "and 80% correctness, and reduced questionnaire completion from "
        "30–40 hours to under 40 seconds. The complete risk-review workflow, "
        "including validation, finishes in under 10 minutes.",
    ),
    (
        "Offer Letter Quality Analysis — Agentic Execution",
        "Technologies: Llama 4 Maverick, multimodal LLMs, prompt-engineered "
        "validation workflows, Meta infrastructure.",
        "Architected and delivered an AI-powered quality-assurance platform "
        "that reviews offer letters before release. A multimodal Visual QA "
        "workflow checks layout, branding, spacing, and formatting, while Text "
        "QA validates candidate attributes, role information, compensation, "
        "dates, and legal clauses against expected data. The platform explains "
        "detected discrepancies and produces an approval or review "
        "recommendation, reducing manual effort and preventing recruiting "
        "document errors at scale.",
    ),
]

STRAIVE_PROJECT = (
    "Obituary Lead Generation",
    "Technologies: GPT-4o, Generative AI, information extraction, lead "
    "qualification and reporting.",
    "Designed and developed a GPT-powered lead-generation tool that discovers "
    "and processes obituary pages, extracts structured information about the "
    "deceased and related family or contact details, and evaluates potential "
    "service needs. GPT-4o converts unstructured page content into qualified "
    "leads for relevant services such as funeral support, estate and mortgage "
    "settlement, and elderly assistance. The workflow deduplicates and enriches "
    "records, produces an explainable lead summary, and automates reporting and "
    "handoff to marketing teams for appropriate outreach.",
)

WIPRO_EXPERIENCE = [
    "Architecting and scaling agentic AI systems for Meta enterprise workflows "
    "using Nexus Runbooks, Agentic Workflow Builder, and Dumont, enabling "
    "autonomous execution, orchestration, and large-scale decision-making.",
    "Designing multi-agent architectures across Meta’s ecosystem, including "
    "Claude-based agents, Eva, and DevMate, with tool integration, memory, and "
    "reasoning for complex workflow and support automation.",
    "Contributing as part of the delivery leadership group for end-to-end "
    "GenAI initiatives, influencing technical direction, use-case "
    "prioritization, stakeholder alignment, and production delivery without "
    "holding direct people-management responsibility.",
]

GRAMENER_LEADERSHIP = (
    "Led a five-member delivery team across Generative AI initiatives, "
    "coordinating solution design, work allocation, execution, and delivery "
    "while remaining hands-on with architecture and implementation."
)

LEADERSHIP_PROFILE = [
    "Leadership experience spans Proxima Systems, Gramener/Straive, and Wipro, "
    "combining direct team management with technical leadership, delivery "
    "ownership, stakeholder influence, and hands-on AI architecture.",
    "At Gramener/Straive, led a five-member team delivering Generative AI "
    "solutions, coordinating priorities, solution design, execution, client "
    "communication, and end-to-end delivery.",
    "At Proxima Systems, led a cross-functional AI/ML team through requirements, "
    "architecture, development, mentoring, deployment, and enterprise "
    "integration.",
    "At Wipro, contributes as part of the leadership group for Meta-focused "
    "GenAI initiatives. Although the role does not include direct people "
    "management, it involves technical direction, architecture ownership, "
    "use-case prioritization, stakeholder alignment, and influencing delivery "
    "decisions across product and engineering teams.",
    "Leadership strengths include translating business problems into AI "
    "roadmaps, creating clarity across cross-functional teams, balancing "
    "hands-on engineering with delivery governance, mentoring practitioners, "
    "and communicating effectively with clients and senior stakeholders.",
]

LEADERSHIP_FAQ_ANSWER = (
    "Yes. His leadership experience spans multiple roles. At Gramener/Straive, "
    "he directly led a five-member Generative AI delivery team. At Proxima "
    "Systems, he led a cross-functional AI/ML team through the complete "
    "delivery lifecycle. At Wipro, he is part of the leadership group for "
    "Meta-focused GenAI initiatives, contributing technical direction, "
    "architecture ownership, prioritization, and stakeholder alignment without "
    "direct people-management responsibility. He has also led larger delivery "
    "and roadmap activities involving 10+ members earlier in his career."
)

PROJECTS_FAQ_ANSWER = (
    "His strongest recent corporate work includes Meta’s People Data Review "
    "agentic workflow, which automates 160+ compliance questions with 99% "
    "groundedness; AWB Runbook Automation using multi-agent orchestration and "
    "MCP-compatible tools; and multimodal Offer Letter Quality Analysis. Other "
    "notable work includes the Axis Max Life agentic SQL chatbot, an OCR and "
    "GenAI legal-document pipeline for LexisNexis-related risk workflows, "
    "Logitech Sherlock text-to-SQL enhancement, HR document assistants, and "
    "financial-risk prediction products at HighRadius."
)


def add_project_before(anchor, project: tuple[str, str, str]) -> None:
    title, technologies, description = project
    anchor.insert_paragraph_before(title, style="Heading 3")
    tech = anchor.insert_paragraph_before(technologies, style="Normal")
    if tech.runs:
        tech.runs[0].italic = True
    anchor.insert_paragraph_before(description, style="Normal")


def find_heading(document: Document, text: str, style: str | None = None):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() != text:
            continue
        if style is None or paragraph.style.name == style:
            return paragraph
    raise RuntimeError(f"Could not find heading: {text}")


def update_project(
    document: Document,
    project: tuple[str, str, str],
    previous_title: str | None = None,
) -> None:
    title, technologies, description = project
    accepted_titles = {title}
    if previous_title:
        accepted_titles.add(previous_title)
    paragraphs = document.paragraphs
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() not in accepted_titles:
            continue
        if paragraph.style.name != "Heading 3":
            continue
        paragraph.text = title
        if index + 1 < len(paragraphs):
            paragraphs[index + 1].text = technologies
            if paragraphs[index + 1].runs:
                paragraphs[index + 1].runs[0].italic = True
        if index + 2 < len(paragraphs):
            paragraphs[index + 2].text = description
        return
    raise RuntimeError(f"Could not update project: {title}")


def main() -> None:
    if not DOCUMENT_PATH.exists():
        raise RuntimeError("Document.docx is missing.")
    if not BACKUP_PATH.exists():
        shutil.copy2(DOCUMENT_PATH, BACKUP_PATH)

    document = Document(DOCUMENT_PATH)
    existing = {paragraph.text.strip() for paragraph in document.paragraphs}

    if WIPRO_PROJECTS[0][0] not in existing:
        straive = find_heading(document, "Straive", "Heading 2")
        straive.insert_paragraph_before("Wipro — Meta", style="Heading 2")
        for project in WIPRO_PROJECTS:
            add_project_before(straive, project)

    if (
        STRAIVE_PROJECT[0] not in existing
        and "Lead Generation from Obituary Websites" not in existing
    ):
        proxima = find_heading(document, "Proxima Systems", "Heading 2")
        add_project_before(proxima, STRAIVE_PROJECT)

    for project in WIPRO_PROJECTS:
        update_project(document, project)
    update_project(
        document,
        STRAIVE_PROJECT,
        previous_title="Lead Generation from Obituary Websites",
    )

    if "CAC Optimizer — Deep Optimization Approach" not in existing:
        technical = find_heading(document, "Technical Capabilities", "Heading 1")
        technical.insert_paragraph_before(
            "CAC Optimizer — Deep Optimization Approach: predictive optimization "
            "using machine-learning algorithms for customer-acquisition-cost "
            "analysis and decision support.",
            style="List Bullet",
        )

    for paragraph in document.paragraphs:
        if paragraph.text.strip() == "Payment Data Prediction (PDP)":
            paragraph.text = "Payment Date Prediction (PDP)"
        if paragraph.text.strip().startswith(
            "Notable work includes an agentic SQL chatbot for Axis Max Life"
        ):
            paragraph.text = PROJECTS_FAQ_ANSWER
        if paragraph.text.strip().startswith(
            "Yes. He has led cross-functional teams of 10+ members"
        ):
            paragraph.text = LEADERSHIP_FAQ_ANSWER

    gramener = find_heading(
        document, "Senior Data Scientist – Gramener – A Straive Company", "Heading 2"
    )
    paragraphs = document.paragraphs
    gramener_position = next(
        index for index, paragraph in enumerate(paragraphs)
        if paragraph._p is gramener._p
    )
    for paragraph in paragraphs[gramener_position + 1 :]:
        if paragraph.style.name.startswith("Heading"):
            break
        if paragraph.style.name.startswith("List Bullet"):
            paragraph.text = GRAMENER_LEADERSHIP
            break

    if "Leadership Profile" not in {
        paragraph.text.strip() for paragraph in document.paragraphs
    }:
        projects = find_heading(document, "Selected Professional Projects", "Heading 1")
        projects.insert_paragraph_before("Leadership Profile", style="Heading 1")
        for statement in LEADERSHIP_PROFILE:
            projects.insert_paragraph_before(statement, style="List Bullet")
    else:
        leadership = find_heading(document, "Leadership Profile", "Heading 1")
        paragraphs = document.paragraphs
        leadership_position = next(
            index for index, paragraph in enumerate(paragraphs)
            if paragraph._p is leadership._p
        )
        bullets = []
        for paragraph in paragraphs[leadership_position + 1 :]:
            if paragraph.style.name.startswith("Heading"):
                break
            if paragraph.style.name.startswith("List Bullet"):
                bullets.append(paragraph)
        for paragraph, statement in zip(bullets, LEADERSHIP_PROFILE):
            paragraph.text = statement

    paragraphs = document.paragraphs
    position = next(
        index
        for index, paragraph in enumerate(paragraphs)
        if paragraph.text.strip() == "Lead Data Scientist – Wipro — Client - Meta"
        and paragraph.style.name == "Heading 2"
    )
    role_paragraphs = []
    for paragraph in paragraphs[position + 1 :]:
        if paragraph.style.name.startswith("Heading"):
            break
        if paragraph.style.name.startswith("List Bullet"):
            role_paragraphs.append(paragraph)
    for paragraph, replacement in zip(role_paragraphs, WIPRO_EXPERIENCE):
        paragraph.text = replacement

    document.save(DOCUMENT_PATH)
    print(f"Updated {DOCUMENT_PATH.name}")
    print(f"Backup: {BACKUP_PATH.name}")


if __name__ == "__main__":
    main()
