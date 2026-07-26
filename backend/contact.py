from __future__ import annotations

import re
from functools import lru_cache
from pathlib import Path
from typing import Iterable

from docx import Document


PHONE_REQUEST_PATTERN = re.compile(
    r"\b(phone(?:\s+number)?|mobile(?:\s+number)?|contact\s+number|"
    r"whatsapp(?:\s+number)?|number\s+to\s+(?:call|contact|reach))\b",
    re.IGNORECASE,
)
PHONE_FOLLOW_UP_PATTERN = re.compile(
    r"\b(present|look|check|recheck|again|find|document|resume|résumé|there)\b",
    re.IGNORECASE,
)
INDIAN_PHONE_PATTERN = re.compile(
    r"(?<!\d)(?:\+?91[\s.-]*)?([6-9](?:[\s.-]*\d){9})(?!\d)"
)
CONTACT_LABEL_PATTERN = re.compile(
    r"\b(phone|mobile|contact|call|whatsapp)\b", re.IGNORECASE
)


def _document_lines(document_path: Path) -> list[str]:
    document = Document(document_path)
    lines = [paragraph.text.strip() for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text.strip() for cell in row.cells))
    return [line for line in lines if line]


@lru_cache(maxsize=8)
def extract_verified_phone(document_path: Path) -> str | None:
    """Return the labelled Indian mobile number from the verified document."""
    lines = _document_lines(document_path)
    labelled_lines = [line for line in lines if CONTACT_LABEL_PATTERN.search(line)]

    for line in [*labelled_lines, *lines]:
        match = INDIAN_PHONE_PATTERN.search(line)
        if not match:
            continue
        local_number = re.sub(r"\D", "", match.group(1))
        return f"+91 {local_number}"
    return None


def is_phone_request(
    query: str,
    conversation: Iterable[tuple[str, str]] = (),
) -> bool:
    if PHONE_REQUEST_PATTERN.search(query):
        return True
    if not PHONE_FOLLOW_UP_PATTERN.search(query):
        return False

    previous_user_messages = [
        content for role, content in conversation if role == "user"
    ][-3:]
    return any(PHONE_REQUEST_PATTERN.search(message) for message in previous_user_messages)
