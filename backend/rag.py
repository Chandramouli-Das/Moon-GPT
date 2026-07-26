from __future__ import annotations

import hashlib
import json
import math
import re
import threading
from collections import Counter
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Callable

import docx
import faiss
import numpy as np

TOKEN_PATTERN = re.compile(r"[a-z0-9+#.]+")
FOLLOW_UP_PATTERN = re.compile(
    r"\b(it|that|this|those|these|them|he|his|more|also|what about|how about)\b",
    re.IGNORECASE,
)
STOP_WORDS = {
    "a", "an", "and", "are", "as", "at", "be", "by", "for", "from", "has",
    "have", "he", "his", "i", "in", "is", "it", "me", "more", "of", "on",
    "or", "that", "the", "this", "to", "was", "what", "when", "where", "which",
    "who", "with", "you",
}
QUERY_ALIASES = {
    "notice": ("notice period", "joining availability", "60 days", "30 days"),
    "join": ("joining availability", "notice period"),
    "salary": ("compensation", "current ctc", "expected ctc"),
    "location": ("Bangalore", "relocation", "work arrangement"),
    "leadership": ("team members led", "mentoring", "stakeholder management"),
    "genai": ("generative ai", "llm", "rag", "agentic ai"),
    "projects": ("project portfolio", "business impact", "production"),
    "experience": ("professional experience", "career timeline"),
    "phone": ("mobile contact phone number call WhatsApp",),
    "mobile": ("phone number contact call WhatsApp",),
    "contact": ("mobile phone number appointment LinkedIn",),
}
WORK_QUERY_PATTERN = re.compile(
    r"\b(work|career|experience|employment|job|role|company|corporate|projects?)\b",
    re.IGNORECASE,
)
FREELANCE_QUERY_PATTERN = re.compile(
    r"\b(freelanc\w*|independent consulting|side projects?|personal projects?)\b",
    re.IGNORECASE,
)
CONTACT_QUERY_PATTERN = re.compile(
    r"\b(phone|mobile|contact|call|whatsapp|reach him|reach chandramouli)\b",
    re.IGNORECASE,
)


@dataclass(frozen=True)
class KnowledgeChunk:
    chunk_id: str
    section: str
    subsection: str
    text: str

    @property
    def label(self) -> str:
        if self.subsection and self.subsection != self.section:
            return f"{self.section} › {self.subsection}"
        return self.section

    @property
    def embedding_text(self) -> str:
        return f"{self.label}\n{self.text}"


@dataclass(frozen=True)
class SearchResult:
    chunk: KnowledgeChunk
    score: float
    dense_score: float
    lexical_score: float


def tokenize(text: str) -> list[str]:
    return [
        token for token in TOKEN_PATTERN.findall(text.lower())
        if token not in STOP_WORDS and len(token) > 1
    ]


def _split_words(text: str, target_words: int, overlap_words: int) -> list[str]:
    words = text.split()
    if len(words) <= target_words:
        return [text.strip()] if text.strip() else []
    step = max(target_words - overlap_words, 1)
    return [
        " ".join(words[start : start + target_words])
        for start in range(0, len(words), step)
        if words[start : start + target_words]
    ]


def load_structured_docx(
    path: Path, target_words: int = 220, overlap_words: int = 35
) -> list[KnowledgeChunk]:
    if not path.exists():
        raise RuntimeError(f"Resume source not found: {path.name}")

    document = docx.Document(path)
    section = "Profile"
    subsection = "Profile"
    grouped: list[tuple[str, str, list[str]]] = []
    current: list[str] = []

    def flush() -> None:
        nonlocal current
        if current:
            grouped.append((section, subsection, current))
            current = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = (paragraph.style.name if paragraph.style else "").lower()
        if style.startswith("heading"):
            flush()
            level_match = re.search(r"(\d+)", style)
            level = int(level_match.group(1)) if level_match else 2
            if level == 1:
                section = text
                subsection = text
            else:
                subsection = text
            continue
        current.append(text)
    flush()

    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" — ".join(cell for cell in cells if cell))
        if rows:
            grouped.append(("Recruiter Quick Reference", "Key facts", rows))

    chunks: list[KnowledgeChunk] = []
    for group_number, (group_section, group_subsection, paragraphs) in enumerate(grouped):
        content = "\n".join(paragraphs)
        for part_number, part in enumerate(
            _split_words(content, target_words, overlap_words)
        ):
            digest = hashlib.sha1(
                f"{group_section}|{group_subsection}|{group_number}|{part_number}|{part}".encode()
            ).hexdigest()[:12]
            chunks.append(
                KnowledgeChunk(
                    chunk_id=digest,
                    section=group_section,
                    subsection=group_subsection,
                    text=part,
                )
            )
    if not chunks:
        raise RuntimeError("The resume source is empty.")
    return chunks


def build_retrieval_query(
    query: str, conversation: list[tuple[str, str]], max_previous_users: int = 2
) -> str:
    cleaned = query.strip()
    if len(tokenize(cleaned)) >= 6 and not FOLLOW_UP_PATTERN.search(cleaned):
        return cleaned
    previous = [
        content.strip()
        for role, content in conversation[:-1]
        if role == "user" and content.strip()
    ][-max_previous_users:]
    return " | ".join([*previous, cleaned]) if previous else cleaned


class RAGEngine:
    def __init__(
        self,
        document_path: Path,
        embedding_model: str,
        embedder: Callable[[list[str]], np.ndarray],
        cache_dir: Path,
    ) -> None:
        self.document_path = document_path
        self.embedding_model = embedding_model
        self.embedder = embedder
        self.cache_dir = cache_dir
        self._lock = threading.Lock()
        self._index: faiss.Index | None = None
        self._chunks: list[KnowledgeChunk] = []
        self._tokens: list[list[str]] = []
        self._document_frequencies: Counter[str] = Counter()
        self._average_length = 1.0
        self.loaded_from_cache = False

    @property
    def ready(self) -> bool:
        return self._index is not None

    @property
    def chunk_count(self) -> int:
        return len(self._chunks)

    def _fingerprint(self) -> str:
        digest = hashlib.sha256()
        digest.update(self.document_path.read_bytes())
        digest.update(self.embedding_model.encode())
        digest.update(b"heading-aware-v2:220:35")
        return digest.hexdigest()

    def _prepare_lexical_index(self) -> None:
        self._tokens = [tokenize(chunk.embedding_text) for chunk in self._chunks]
        self._document_frequencies = Counter()
        for tokens in self._tokens:
            self._document_frequencies.update(set(tokens))
        self._average_length = (
            sum(map(len, self._tokens)) / len(self._tokens) if self._tokens else 1.0
        )

    def _load_cache(self, fingerprint: str) -> bool:
        metadata_path = self.cache_dir / "metadata.json"
        index_path = self.cache_dir / "index.faiss"
        if not metadata_path.exists() or not index_path.exists():
            return False
        try:
            metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
            if metadata.get("fingerprint") != fingerprint:
                return False
            self._chunks = [
                KnowledgeChunk(**chunk) for chunk in metadata.get("chunks", [])
            ]
            self._index = faiss.read_index(str(index_path))
            if not self._chunks or self._index.ntotal != len(self._chunks):
                return False
            self._prepare_lexical_index()
            self.loaded_from_cache = True
            return True
        except (OSError, ValueError, TypeError, json.JSONDecodeError):
            return False

    def _save_cache(self, fingerprint: str) -> None:
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        metadata_path = self.cache_dir / "metadata.json"
        index_path = self.cache_dir / "index.faiss"
        metadata = {
            "fingerprint": fingerprint,
            "embedding_model": self.embedding_model,
            "chunks": [asdict(chunk) for chunk in self._chunks],
        }
        metadata_path.write_text(
            json.dumps(metadata, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        faiss.write_index(self._index, str(index_path))

    def ensure_ready(self) -> None:
        if self._index is not None:
            return
        with self._lock:
            if self._index is not None:
                return
            fingerprint = self._fingerprint()
            if self._load_cache(fingerprint):
                return
            self._chunks = load_structured_docx(self.document_path)
            matrix = self.embedder(
                [chunk.embedding_text for chunk in self._chunks]
            ).astype(np.float32)
            faiss.normalize_L2(matrix)
            self._index = faiss.IndexFlatIP(matrix.shape[1])
            self._index.add(matrix)
            self._prepare_lexical_index()
            self.loaded_from_cache = False
            self._save_cache(fingerprint)

    def _lexical_scores(self, query_tokens: list[str]) -> np.ndarray:
        scores = np.zeros(len(self._chunks), dtype=np.float32)
        if not query_tokens:
            return scores
        total = max(len(self._chunks), 1)
        k1, b = 1.5, 0.75
        for index, tokens in enumerate(self._tokens):
            counts = Counter(tokens)
            length_ratio = len(tokens) / self._average_length
            score = 0.0
            for token in query_tokens:
                frequency = counts[token]
                if not frequency:
                    continue
                document_frequency = self._document_frequencies[token]
                inverse_frequency = math.log(
                    1 + (total - document_frequency + 0.5) / (document_frequency + 0.5)
                )
                score += inverse_frequency * (
                    frequency * (k1 + 1)
                    / (frequency + k1 * (1 - b + b * length_ratio))
                )
            scores[index] = score
        maximum = float(scores.max()) if len(scores) else 0.0
        return scores / maximum if maximum > 0 else scores

    @staticmethod
    def _expanded_query(query: str) -> str:
        additions: list[str] = []
        lowered = query.lower()
        for key, aliases in QUERY_ALIASES.items():
            if key in lowered:
                additions.extend(aliases)
        return f"{query} {' '.join(additions)}".strip()

    def search(self, query: str, count: int = 5) -> list[SearchResult]:
        self.ensure_ready()
        assert self._index is not None
        expanded = self._expanded_query(query)
        query_vector = self.embedder([expanded]).astype(np.float32)
        faiss.normalize_L2(query_vector)
        candidate_count = min(max(count * 4, 12), len(self._chunks))
        dense_scores, positions = self._index.search(query_vector, candidate_count)
        dense_by_position = {
            int(position): max(float(score), 0.0)
            for score, position in zip(dense_scores[0], positions[0])
            if position >= 0
        }
        lexical = self._lexical_scores(tokenize(expanded))
        candidates = set(dense_by_position)
        candidates.update(np.argsort(lexical)[-candidate_count:].tolist())
        prefer_corporate = bool(WORK_QUERY_PATTERN.search(query)) and not bool(
            FREELANCE_QUERY_PATTERN.search(query)
        )
        prefer_contact = bool(CONTACT_QUERY_PATTERN.search(query))

        ranked = sorted(
            (
                SearchResult(
                    chunk=self._chunks[position],
                    # Generic work questions should lead with verified employment
                    # and employer-owned delivery, not personal/side projects.
                    score=0.68 * dense_by_position.get(position, 0.0)
                    + 0.32 * float(lexical[position])
                    + (
                        0.14
                        if prefer_corporate
                        and self._chunks[position].section == "Professional Experience"
                        else 0.08
                        if prefer_corporate
                        and self._chunks[position].section
                        == "Selected Professional Projects"
                        else -0.10
                        if prefer_corporate
                        and self._chunks[position].section
                        == "Personal AI & Data Projects"
                        else 0.0
                    )
                    + (
                        0.65
                        if prefer_contact
                        and self._chunks[position].section
                        in {"Contact & Professional Links", "Recruiter Quick Reference"}
                        else 0.0
                    ),
                    dense_score=dense_by_position.get(position, 0.0),
                    lexical_score=float(lexical[position]),
                )
                for position in candidates
            ),
            key=lambda result: result.score,
            reverse=True,
        )

        selected: list[SearchResult] = []
        section_counts: Counter[str] = Counter()
        for result in ranked:
            if section_counts[result.chunk.label] >= 2:
                continue
            selected.append(result)
            section_counts[result.chunk.label] += 1
            if len(selected) == count:
                break
        return selected
